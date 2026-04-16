from docx import Document


def validate_docx(docx_path):
    try:
        Document(docx_path)
    except Exception:
        return False

    return True


def docx_to_text(docx_path):
    if not validate_docx(docx_path):
        return False

    try:
        doc = Document(docx_path)
        full_text = []

        for para in doc.paragraphs:
            full_text.append(para.text)

        return '\n'.join(full_text)

    except Exception:
        return False